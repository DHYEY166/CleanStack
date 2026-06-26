import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "package"))

try:
    import sentry_sdk
    from sentry_sdk.integrations.aws_lambda import AwsLambdaIntegration
    sentry_sdk.init(
        dsn=os.environ.get("SENTRY_DSN", ""),
        integrations=[AwsLambdaIntegration(timeout_warning=True)],
        traces_sample_rate=0.1,
    )
except ImportError:
    pass

import json
import io
import re
import boto3
import psycopg2
import requests
import pandas as pd
import numpy as np
from collections import Counter

s3 = boto3.client("s3")
secrets = boto3.client("secretsmanager")

DOCUMENT_EXTENSIONS = {"pdf", "docx", "doc"}

SENTINEL_VALUES = {
    # Explicit null markers
    "", "n/a", "na", "null", "none", "unknown", "undefined", "not available",
    "not applicable", "not provided", "not specified", "not given",
    # Punctuation sentinels
    "-", "--", "---", "----", ".", "..", "...",
    "?", "??", "???", "#", "##",
    # Coded sentinels
    "0", "00", "000", "-1", "99", "999", "9999", "99999", "-99", "-999",
    # Boolean-as-sentinel
    "false", "nil", "nan", "missing", "void",
    # State sentinels
    "n.a.", "n.a", "#n/a", "#null!", "tbd", "tbc", "pending", "not set",
    "to be determined", "to be confirmed", "unknown value",
    # Excel/CSV export artifacts
    "#value!", "#ref!", "#div/0!", "#name?", "#num!", "#error!",
    "error", "err", "null value", "blank", "empty",
    # Common filler
    "x", "xx", "xxx", "test", "temp", "placeholder", "sample",
    # Numeric as string
    "0.0", "0.00", "-1.0", "inf", "-inf",
}

DOMAIN_KEYWORDS = {
    "contract":  ["agreement", "clause", "party", "whereas", "hereinafter", "indemnify", "termination", "obligations"],
    "medical":   ["patient", "diagnosis", "prescription", "physician", "clinical", "dosage", "treatment", "symptoms"],
    "hr":        ["employee", "salary", "compensation", "performance", "payroll", "benefits", "onboarding", "recruiter"],
    "invoice":   ["invoice", "billing", "payment", "amount due", "vendor", "purchase order", "remittance", "net 30"],
    "legal":     ["plaintiff", "defendant", "court", "jurisdiction", "liability", "statute", "affidavit", "counsel"],
    "financial": ["revenue", "ebitda", "balance sheet", "fiscal", "quarterly", "dividend", "earnings", "amortization"],
}


def get_db_conn():
    from urllib.parse import urlparse
    url = os.environ["DATABASE_URL"]
    p = urlparse(url)
    host, port, user, dbname = p.hostname, p.port or 5432, p.username, p.path.lstrip("/")
    rds = boto3.client("rds", region_name=os.environ.get("AWS_REGION", "us-east-1"))
    token = rds.generate_db_auth_token(DBHostname=host, Port=port, DBUsername=user)
    return psycopg2.connect(host=host, port=port, user=user, password=token, dbname=dbname, sslmode="require")


def detect_encoding(file_bytes: bytes) -> str:
    try:
        import chardet
        result = chardet.detect(file_bytes[:8192])
        enc = result.get("encoding") or "utf-8"
        confidence = result.get("confidence", 0.0)
        return enc if confidence > 0.7 else "utf-8"
    except ImportError:
        return "utf-8"


def detect_format(key: str, content_type: str, file_bytes: bytes = b"") -> str:
    if file_bytes:
        magic = file_bytes[:8]
        if magic[:4] == b'\xd0\xcf\x11\xe0':  # OLE2 compound = old .xls
            return "xls"
        if magic[:4] == b'PK\x03\x04':  # ZIP-based = .xlsx or .docx
            ext = key.rsplit(".", 1)[-1].lower()
            return ext if ext in ("xlsx", "docx") else "xlsx"
        if magic[:4] == b'%PDF':
            return "pdf"
    return key.rsplit(".", 1)[-1].lower()


def extract_text(file_bytes: bytes, fmt: str) -> str:
    if fmt == "pdf":
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            parts.append(" | ".join(str(c) if c else "" for c in row))
                else:
                    text = page.extract_text()
                    if text:
                        parts.append(text)
        return "\n\n".join(parts)
    elif fmt == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        # Main body paragraphs
        parts.extend(p.text for p in doc.paragraphs)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        # Headers and footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                parts.append(para.text)
            for para in section.footer.paragraphs:
                parts.append(para.text)
        # Text boxes (DrawingML)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for txbx in doc.element.findall('.//w:txbxContent//w:p', ns):
            text = "".join(r.text or "" for r in txbx.findall('.//w:r/w:t', ns))
            if text:
                parts.append(text)
        return "\n".join(p for p in parts if p)
    else:
        return file_bytes.decode("utf-8", errors="replace")


def _val_pattern(val: str) -> str:
    s = re.sub(r'[A-Za-z]+', 'A', val)
    s = re.sub(r'\d+', 'N', s)
    return s


def profile_document(text: str) -> dict:
    lines = text.splitlines()
    words = text.split()
    non_blank_lines = [l for l in lines if l.strip()]

    # — PII counts —
    email_count = len(re.findall(r'[\w.+-]+@[\w-]+\.\w+', text))
    phone_count = len(re.findall(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', text))
    ssn_count   = len(re.findall(r'\b\d{3}-\d{2}-\d{4}\b', text))
    cc_count    = len(re.findall(r'\b(?:\d{4}[- ]?){3}\d{4}\b', text))
    html_tags   = len(re.findall(r'<[^>]+>', text))
    blank_lines = sum(1 for l in lines if not l.strip())

    # — Repeated lines (headers/footers) —
    line_counts = Counter(l.strip() for l in lines if l.strip() and len(l.strip()) < 150)
    repeated = {l: c for l, c in line_counts.items() if c >= 3}
    repeated_examples = dict(list(sorted(repeated.items(), key=lambda x: -x[1])[:5]))

    # — Line length distribution —
    lengths = [len(l) for l in non_blank_lines]
    avg_line_len  = round(sum(lengths) / max(len(lengths), 1), 1)
    short_line_pct = round(sum(1 for l in lengths if l < 40) / max(len(lengths), 1) * 100, 1)

    # — Encoding errors —
    enc_errors = re.findall(r'[â€™â€œÃ©Ã¨Ã®Ã´Ã ï»¿�â]', text)
    enc_error_examples = list({c for c in enc_errors})[:5]

    # — Named entity hints —
    person_matches = list(set(re.findall(r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b', text)))
    org_matches    = list(set(re.findall(
        r'\b[A-Z][A-Za-z\s]{2,30}(?:Inc|LLC|Ltd|LLP|Corp|Co|Company|Group|Holdings|Technologies|Solutions)\b', text
    )))

    # — Domain inference —
    text_lower = text.lower()
    domain_scores = {d: sum(text_lower.count(k) for k in kws) for d, kws in DOMAIN_KEYWORDS.items()}
    best_domain = max(domain_scores, key=domain_scores.get)
    domain_confidence = domain_scores[best_domain]
    inferred_domain = best_domain if domain_confidence >= 2 else "general"

    # — Quality score —
    pii_count      = email_count + phone_count + ssn_count + cc_count
    pii_penalty    = min(pii_count * 5, 40)
    html_penalty   = min(html_tags * 0.5, 20)
    blank_penalty  = min(blank_lines / max(len(lines), 1) * 100 * 0.3, 20)
    ner_penalty    = min(len(person_matches) * 1.5, 15)
    enc_penalty    = min(len(enc_errors) * 0.5, 10)
    header_penalty = min(len(repeated) * 1.0, 10)
    quality_score  = max(0, round(100 - pii_penalty - html_penalty - blank_penalty
                                  - ner_penalty - enc_penalty - header_penalty))

    mid = len(text) // 2
    doc_stats = {
        "word_count":             len(words),
        "char_count":             len(text),
        "blank_line_count":       blank_lines,
        "avg_line_length":        avg_line_len,
        "short_line_pct":         short_line_pct,
        "pii_detected": {
            "emails":       email_count,
            "phones":       phone_count,
            "ssns":         ssn_count,
            "credit_cards": cc_count,
        },
        "html_tag_count":         html_tags,
        "repeated_line_count":    len(repeated),
        "repeated_line_examples": repeated_examples,
        "encoding_error_count":   len(enc_errors),
        "encoding_error_examples": enc_error_examples,
        "person_name_count":      len(person_matches),
        "person_name_examples":   person_matches[:5],
        "org_count":              len(org_matches),
        "org_examples":           org_matches[:3],
        "inferred_domain":        inferred_domain,
        "domain_confidence":      domain_confidence,
        "sample_text":            text[:2000],
        "mid_sample_text":        text[mid: mid + 500],
    }

    return {
        "quality_score":        quality_score,
        "total_rows":           len(lines),
        "null_percentage":      0.0,
        "duplicate_percentage": 0.0,
        "type_mismatch_count":  0,
        "outlier_count":        0,
        "column_stats":         doc_stats,
    }


def load_dataframe(file_bytes: bytes, fmt: str) -> pd.DataFrame:
    buf = io.BytesIO(file_bytes)

    if fmt == "csv":
        encoding = detect_encoding(file_bytes)
        sample = file_bytes[:4096].decode(encoding, errors="replace")
        sep = "\t" if sample.count("\t") > sample.count(",") else ","
        return pd.read_csv(
            io.BytesIO(file_bytes), sep=sep,
            dtype=str, keep_default_na=False, low_memory=False,
            encoding=encoding, encoding_errors="replace",
        )
    elif fmt == "txt":
        encoding = detect_encoding(file_bytes)
        sample = file_bytes[:4096].decode(encoding, errors="replace")
        counts = {s: sample.count(s) for s in [",", "\t", "|", ";"]}
        sep = max(counts, key=counts.get)
        if counts[sep] < 2:
            return pd.read_csv(
                io.BytesIO(file_bytes), sep=r'\s+',
                dtype=str, keep_default_na=False, engine='python',
                encoding=encoding, encoding_errors="replace",
            )
        return pd.read_csv(
            io.BytesIO(file_bytes), sep=sep,
            dtype=str, keep_default_na=False, low_memory=False,
            encoding=encoding, encoding_errors="replace",
        )
    elif fmt == "tsv":
        encoding = detect_encoding(file_bytes)
        return pd.read_csv(
            buf, sep="\t",
            dtype=str, keep_default_na=False, low_memory=False,
            encoding=encoding, encoding_errors="replace",
        )
    elif fmt in ("json", "jsonl"):
        text = file_bytes.decode("utf-8", errors="replace").strip()
        if fmt == "jsonl":
            # Strip comment lines before parsing
            lines = [l for l in text.splitlines() if not l.strip().startswith("//")]
            text_clean = "\n".join(lines)
            try:
                return pd.read_json(io.BytesIO(text_clean.encode()), lines=True)
            except Exception:
                pass
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                return pd.json_normalize(parsed)
            elif isinstance(parsed, dict):
                for v in parsed.values():
                    if isinstance(v, list):
                        return pd.json_normalize(v)
                return pd.json_normalize([parsed])
        except Exception:
            pass
        try:
            return pd.read_json(io.BytesIO(file_bytes), lines=True)
        except Exception:
            return pd.read_json(io.BytesIO(file_bytes))
    elif fmt in ("xlsx", "xls"):
        xl = pd.ExcelFile(buf)
        df = xl.parse(xl.sheet_names[0], dtype=str, keep_default_na=False)
        # Forward-fill merged cells (NaN after merge top-left = merged cell artifact)
        df = df.ffill(axis=0)
        return df
    elif fmt == "xml":
        from lxml import etree
        root = etree.fromstring(file_bytes)
        rows = [{child.tag: child.text for child in elem} for elem in root]
        if not rows:
            rows = [{root.tag: root.text}]
        return pd.DataFrame(rows)
    else:
        raise ValueError(f"Unsupported format: {fmt}")


def compute_quality_score(df: pd.DataFrame) -> dict:
    total_cells = df.size or 1
    total_rows  = len(df)
    total_cols  = len(df.columns)

    null_count = df.isnull().sum().sum()
    null_pct   = round(null_count / total_cells * 100, 2)

    dup_count = df.duplicated().sum()
    dup_pct   = round(dup_count / max(total_rows, 1) * 100, 2)

    type_mismatches = 0
    for col in df.columns:
        if df[col].dtype == object:
            numeric_count = pd.to_numeric(df[col], errors="coerce").notna().sum()
            if 0 < numeric_count < len(df[col]):
                type_mismatches += 1

    # dtype=str: coerce object cols to numeric before outlier detection
    outlier_count = 0
    for col in df.columns:
        _num = pd.to_numeric(df[col], errors="coerce") if df[col].dtype == object else df[col]
        if _num.notna().sum() < max(len(_num) * 0.5, 2):
            continue
        q1, q3 = _num.quantile(0.25), _num.quantile(0.75)
        iqr = q3 - q1
        if iqr > 0:
            outlier_count += int(((_num < q1 - 1.5 * iqr) | (_num > q3 + 1.5 * iqr)).sum())

    # Sentinel and whitespace — dataset-level aggregates
    total_sentinel_count = 0
    whitespace_padded_cols = 0

    column_stats = {}
    for col in df.columns:
        series = df[col]
        n = len(series)

        col_stat: dict = {
            "type":         str(series.dtype),
            "null_count":   int(series.isnull().sum()),
            "null_pct":     round(series.isnull().mean() * 100, 2),
            "unique_count": int(series.nunique()),
            "sample_values": [
                str(v) if isinstance(v, (int, float)) and abs(v) > 1e15 else v
                for v in series.dropna().head(20).tolist()
            ],
        }

        # dtype=str: try numeric coercion for min/max/outlier stats
        _num_series = series if pd.api.types.is_numeric_dtype(series) else pd.to_numeric(series, errors="coerce")
        if _num_series.notna().sum() >= max(len(_num_series) * 0.5, 2):
            col_stat["min"] = float(_num_series.min()) if not _num_series.empty else None
            col_stat["max"] = float(_num_series.max()) if not _num_series.empty else None
            q1, q3 = _num_series.quantile(0.25), _num_series.quantile(0.75)
            iqr = q3 - q1
            if iqr > 0:
                outliers = _num_series[(_num_series < q1 - 1.5 * iqr) | (_num_series > q3 + 1.5 * iqr)]
                col_stat["outlier_examples"] = [float(v) for v in outliers.head(3).tolist()]

        if series.dtype == object:
            str_series = series.astype(str).str.strip().str.lower()

            # Sentinel detection
            sentinel_count = int(str_series.isin(SENTINEL_VALUES).sum())
            col_stat["sentinel_count"] = sentinel_count
            col_stat["sentinel_pct"]   = round(sentinel_count / max(n, 1) * 100, 2)
            col_stat["true_null_pct"]  = round((col_stat["null_count"] + sentinel_count) / max(n, 1) * 100, 2)
            total_sentinel_count += sentinel_count

            # Sentinel examples (distinct values found)
            sentinel_vals_found = series.astype(str).str.strip()[
                series.astype(str).str.strip().str.lower().isin(SENTINEL_VALUES)
            ].unique().tolist()
            col_stat["sentinel_examples"] = [str(v) for v in sentinel_vals_found[:5]]

            # Whitespace-padded count
            raw_str = series.dropna().astype(str)
            padded = int((raw_str != raw_str.str.strip()).sum())
            col_stat["whitespace_padded_count"] = padded
            if padded > 0:
                whitespace_padded_cols += 1

            # String pattern diversity
            patterns = raw_str.apply(_val_pattern).value_counts().head(6)
            col_stat["string_patterns"]       = {str(k): int(v) for k, v in patterns.items()}
            col_stat["distinct_pattern_count"] = int(raw_str.apply(_val_pattern).nunique())

            # Value frequency for low-cardinality columns
            if series.nunique() <= 50:
                top10 = series.value_counts(dropna=False).head(10)
                col_stat["value_counts"] = {str(k): int(v) for k, v in top10.items()}

        column_stats[str(col)] = col_stat

    # Dataset-level sentinel pct
    total_object_cells = int(df.select_dtypes(include="object").size) or 1
    sentinel_pct_overall = round(total_sentinel_count / total_object_cells * 100, 2)

    # Penalties
    null_penalty     = min(null_pct * 0.5, 30)
    dup_penalty      = min(dup_pct * 0.3, 20)
    type_penalty     = min(type_mismatches * 5, 20)
    outlier_penalty  = min(outlier_count / max(total_rows, 1) * 100 * 0.1, 10)
    sentinel_penalty = min(sentinel_pct_overall * 0.4, 15)
    ws_penalty       = min(whitespace_padded_cols / max(total_cols, 1) * 100 * 0.1, 5)
    score = max(0, round(100 - null_penalty - dup_penalty - type_penalty
                         - outlier_penalty - sentinel_penalty - ws_penalty))

    return {
        "quality_score":          score,
        "total_rows":             total_rows,
        "null_percentage":        null_pct,
        "duplicate_percentage":   dup_pct,
        "type_mismatch_count":    type_mismatches,
        "outlier_count":          outlier_count,
        "sentinel_pct_overall":   sentinel_pct_overall,
        "whitespace_padded_cols": whitespace_padded_cols,
        "column_stats":           column_stats,
    }


def detect_signals(df: pd.DataFrame, column_stats: dict) -> dict:
    """Detect dataset signals for gate-controlled rule suggestion in suggest-transforms."""
    signals = {}
    n = len(df)

    # ── ffill candidate ───────────────────────────────────────────────────────
    DATETIME_HINTS = ("date", "time", "created", "updated", "timestamp", "period",
                      "month", "year", "day")
    date_cols = [c for c in df.columns if any(h in str(c).lower() for h in DATETIME_HINTS)]
    ffill_cols = []
    for col in date_cols:
        null_pct = df[col].isnull().mean()
        if 0.01 < null_pct < 0.60:
            ffill_cols.append(col)
    for col in df.select_dtypes(include="object").columns:
        if col in ffill_cols:
            continue
        null_pct = df[col].isnull().mean()
        if 0.01 < null_pct < 0.40 and any(h in str(col).lower() for h in DATETIME_HINTS):
            ffill_cols.append(col)
    signals["has_ffill_candidate"] = len(ffill_cols) > 0
    signals["ffill_candidate_cols"] = ffill_cols
    signals["ffill_confidence"] = min(len(ffill_cols) * 0.3, 1.0)

    # ── outlier signal ────────────────────────────────────────────────────────
    # dtype=str: coerce object cols to numeric before IQR analysis
    outlier_cols = []
    for col in df.columns:
        _num = pd.to_numeric(df[col], errors="coerce") if df[col].dtype == object else df[col]
        if _num.notna().sum() < max(len(_num) * 0.5, 2):
            continue
        q1, q3 = _num.quantile(0.25), _num.quantile(0.75)
        iqr = q3 - q1
        if iqr > 0:
            if ((_num < q1 - 1.5 * iqr) | (_num > q3 + 1.5 * iqr)).sum() > 0:
                outlier_cols.append(col)
    signals["has_outliers"] = len(outlier_cols) > 0
    signals["outlier_cols"] = outlier_cols
    signals["outlier_confidence"] = min(len(outlier_cols) * 0.3, 1.0)

    # ── boolean columns ───────────────────────────────────────────────────────
    BOOL_VALUES = {"true", "false", "yes", "no", "1", "0", "t", "f", "y", "n", "on", "off"}
    bool_cols = []
    for col in df.select_dtypes(include="object").columns:
        vals = set(df[col].dropna().str.lower().unique()) if hasattr(df[col], 'str') else set()
        if vals and vals.issubset(BOOL_VALUES) and len(vals) <= 4:
            bool_cols.append(col)
    signals["has_boolean_column"] = len(bool_cols) > 0
    signals["boolean_cols"] = bool_cols
    signals["boolean_confidence"] = 0.95 if bool_cols else 0.0

    # ── multi-currency ────────────────────────────────────────────────────────
    CURRENCY_SYMBOLS = {"$", "€", "£", "¥", "₹", "₩", "CHF", "CAD", "AUD"}
    currency_cols = []
    all_symbols_found: set = set()
    for col in df.select_dtypes(include="object").columns:
        sample = df[col].dropna().head(100).astype(str)
        symbols = set()
        for sym in CURRENCY_SYMBOLS:
            if sample.str.contains(re.escape(sym), regex=False).any():
                symbols.add(sym)
        if len(symbols) >= 2:
            currency_cols.append(col)
            all_symbols_found.update(symbols)
    signals["has_multi_currency"] = len(currency_cols) > 0
    signals["currency_symbols_found"] = list(all_symbols_found)
    signals["currency_cols"] = currency_cols
    signals["multi_currency_confidence"] = 0.85 if len(all_symbols_found) >= 2 else 0.0

    # ── split column ──────────────────────────────────────────────────────────
    SPLIT_DELIMITERS = ["|", ";", "::", " - ", "/", "\\"]
    split_cols = []
    split_delimiters: dict = {}
    for col in df.select_dtypes(include="object").columns:
        sample = df[col].dropna().head(100).astype(str)
        for delim in SPLIT_DELIMITERS:
            if sample.str.contains(re.escape(delim), regex=False).mean() > 0.70:
                split_cols.append(col)
                split_delimiters[col] = delim
                break
    signals["has_split_candidate"] = len(split_cols) > 0
    signals["split_cols"] = split_cols
    signals["split_delimiters"] = split_delimiters
    signals["split_confidence"] = 0.70 if split_cols else 0.0

    # ── messy column headers ──────────────────────────────────────────────────
    messy = [c for c in df.columns if re.search(r'[A-Z\s\-\.\/\\#@!%]', str(c)) or " " in str(c)]
    signals["has_messy_headers"] = len(messy) > 0
    signals["messy_header_count"] = len(messy)

    # ── scale ─────────────────────────────────────────────────────────────────
    signals["row_count"] = n
    signals["needs_chunked_processing"] = n > 500_000

    # ── numeric locale (European: 1.234,56 format) ────────────────────────────
    locale_cols = []
    for col in df.select_dtypes(include="object").columns:
        sample = df[col].dropna().head(50).astype(str)
        if sample.str.match(r'^\d{1,3}(\.\d{3})+(,\d+)?$').mean() > 0.5:
            locale_cols.append(col)
    signals["has_numeric_locale"] = len(locale_cols) > 0
    signals["locale_cols"] = locale_cols

    return signals


def handler(event, context):
    record = event["Records"][0]["s3"]
    bucket = record["bucket"]["name"]
    key    = record["object"]["key"].replace("+", " ")

    parts = key.split("/")
    if len(parts) < 4:
        print(f"[profiler] Skipping non-run key: {key}")
        return {"statusCode": 200, "body": "skipped"}

    obj        = s3.get_object(Bucket=bucket, Key=key)
    file_bytes = obj["Body"].read()
    # Strip UTF-8 BOM if present (prevents invisible char in first column name)
    if file_bytes[:3] == b'\xef\xbb\xbf':
        file_bytes = file_bytes[3:]
    fmt        = detect_format(key, obj.get("ContentType", ""), file_bytes)
    run_id     = parts[2]

    conn = get_db_conn()
    cur  = conn.cursor()

    try:
        cur.execute(
            "UPDATE pipeline_runs SET status = 'profiling' WHERE id = %s AND status = 'pending' RETURNING id",
            (run_id,)
        )
        claimed = cur.fetchone()
        conn.commit()

        if not claimed:
            print(f"[profiler] Run {run_id} already claimed — skipping duplicate invocation")
            cur.close()
            conn.close()
            return {"statusCode": 200, "run_id": run_id, "skipped": True}

        if fmt in DOCUMENT_EXTENSIONS:
            mode = "document"
        elif fmt == "txt":
            try:
                df_test = load_dataframe(file_bytes, fmt)
                mode = "tabular" if len(df_test.columns) >= 3 else "document"
            except Exception:
                mode = "document"
        else:
            mode = "tabular"

        cur.execute("UPDATE pipeline_runs SET mode = %s WHERE id = %s", (mode, run_id))
        conn.commit()

        class _NpEncoder(json.JSONEncoder):
            def default(self, obj):
                if isinstance(obj, (np.integer,)):  return int(obj)
                if isinstance(obj, (np.floating,)): return float(obj)
                if isinstance(obj, np.ndarray):     return obj.tolist()
                return super().default(obj)

        if mode == "document":
            text = extract_text(file_bytes, fmt)
            text_key = "/".join(key.rsplit("/", 1)[:-1]) + "/extracted_text.txt"
            s3.put_object(Bucket=bucket, Key=text_key, Body=text.encode("utf-8"), ContentType="text/plain")
            profile = profile_document(text)
        else:
            df = load_dataframe(file_bytes, fmt)
            if df.empty or len(df.columns) <= 1:
                raise ValueError(
                    f"File could not be parsed as structured tabular data "
                    f"({len(df.columns)} column(s), {len(df)} row(s) detected). "
                    f"CleanStack requires structured data with multiple columns."
                )
            profile = compute_quality_score(df)
            # Inject signal detection into column_stats (Phase B — no schema change needed)
            signals = detect_signals(df, profile["column_stats"])
            profile["column_stats"]["_signals"] = signals

        cur.execute(
            """INSERT INTO data_profiles
               (run_id, stage, quality_score, total_rows, null_percentage,
                duplicate_percentage, type_mismatch_count, outlier_count, column_stats)
               VALUES (%s, 'raw', %s, %s, %s, %s, %s, %s, %s)""",
            (
                run_id,
                float(profile["quality_score"]),
                int(profile["total_rows"]),
                float(profile["null_percentage"]),
                float(profile["duplicate_percentage"]),
                int(profile["type_mismatch_count"]),
                int(profile["outlier_count"]),
                json.dumps(profile.get("column_stats", {}), cls=_NpEncoder),
            ),
        )

        cur.execute(
            "UPDATE pipeline_runs SET status = 'awaiting_ai', row_count_raw = %s WHERE id = %s",
            (profile["total_rows"], run_id),
        )
        conn.commit()

        app_url        = os.environ["APP_URL"]
        webhook_secret = os.environ["WEBHOOK_SECRET"]
        requests.post(
            f"{app_url}/api/webhooks/profile-complete",
            json={"run_id": run_id},
            headers={"x-webhook-secret": webhook_secret},
            timeout=300,
        )

    except Exception as e:
        conn.rollback()
        cur.execute(
            "UPDATE pipeline_runs SET status = 'failed', error_message = %s WHERE id = %s",
            (str(e), run_id),
        )
        conn.commit()
        raise
    finally:
        cur.close()
        conn.close()

    return {"statusCode": 200, "run_id": run_id}
