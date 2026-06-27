import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "package"))

try:
    import sentry_sdk
    from sentry_sdk.integrations.aws_lambda import AwsLambdaIntegration
    sentry_sdk.init(
        dsn=os.environ.get("SENTRY_DSN", ""),
        integrations=[AwsLambdaIntegration(timeout_warning=True)],
        traces_sample_rate=0.1,
    )
except ImportError:
    pass

import json
import io
import re
import hashlib
import boto3
import psycopg2
import pandas as pd
import numpy as np

DOCUMENT_EXTENSIONS = {"pdf", "docx"}

s3 = boto3.client("s3")
sns = boto3.client("sns")
secrets = boto3.client("secretsmanager")


class _NpEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.integer): return int(obj)
        if isinstance(obj, np.floating): return None if np.isnan(obj) else float(obj)
        if isinstance(obj, np.ndarray): return obj.tolist()
        return super().default(obj)

def _sanitize_nan(obj):
    """Recursively replace float NaN/Inf with None so PostgreSQL JSON accepts it."""
    if isinstance(obj, float) and (obj != obj or obj == float('inf') or obj == float('-inf')):
        return None
    if isinstance(obj, dict):
        return {k: _sanitize_nan(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize_nan(v) for v in obj]
    return obj


def save_dataframe(df: pd.DataFrame, fmt: str) -> tuple[bytes, str, str]:
    """Return (file_bytes, content_type, extension) in native format."""
    if fmt == "csv":
        buf = io.BytesIO()
        df.to_csv(buf, index=False)
        return buf.getvalue(), "text/csv", "csv"

    elif fmt == "txt":
        buf = io.BytesIO()
        df.to_csv(buf, index=False)
        return buf.getvalue(), "text/plain", "txt"

    elif fmt == "tsv":
        buf = io.BytesIO()
        df.to_csv(buf, sep="\t", index=False)
        return buf.getvalue(), "text/tab-separated-values", "tsv"

    elif fmt == "json":
        json_str = df.to_json(orient="records", indent=2, force_ascii=False)
        return (json_str or "[]").encode("utf-8"), "application/json", "json"

    elif fmt == "jsonl":
        json_str = df.to_json(orient="records", lines=True, force_ascii=False)
        return (json_str or "").encode("utf-8"), "application/x-ndjson", "jsonl"

    elif fmt in ("xlsx", "xls"):
        buf = io.BytesIO()
        df.to_excel(buf, index=False, engine="openpyxl")
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"

    elif fmt == "xml":
        from lxml import etree
        root = etree.Element("records")
        for _, row in df.iterrows():
            record = etree.SubElement(root, "record")
            for col, val in row.items():
                child = etree.SubElement(record, str(col).replace(" ", "_"))
                child.text = "" if (val is None or (isinstance(val, float) and np.isnan(val))) else str(val)
        xml_bytes = etree.tostring(root, pretty_print=True, xml_declaration=True, encoding="UTF-8")
        return xml_bytes, "application/xml", "xml"

    elif fmt == "parquet":
        buf = io.BytesIO()
        df.to_parquet(buf, index=False, engine="pyarrow")
        return buf.getvalue(), "application/octet-stream", "parquet"

    else:
        buf = io.BytesIO()
        df.to_csv(buf, index=False)
        return buf.getvalue(), "text/csv", "csv"


def extract_text(file_bytes: bytes, fmt: str) -> str:
    if fmt == "pdf":
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n\n".join(pages)
    elif fmt == "docx":
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        return file_bytes.decode("utf-8", errors="replace")


def apply_transforms_pdf(file_bytes: bytes, rules: list[dict]) -> tuple[bytes, str, str]:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")

    for page in doc:
        page_text = page.get_text("text")

        for rule in rules:
            rtype = rule["rule_type"]
            params = _parse_params_simple(rule.get("parameters", {}))
            try:
                if rtype == "strip_pii":
                    patterns = [
                        r'[\w.+-]+@[\w-]+\.\w+',
                        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
                        r'\b\d{3}-\d{2}-\d{4}\b',
                        r'\b(?:\d{4}[- ]?){3}\d{4}\b',
                    ]
                    replacements = ["[EMAIL REDACTED]", "[PHONE REDACTED]", "[SSN REDACTED]", "[CC REDACTED]"]
                    for pat, repl in zip(patterns, replacements):
                        for match in set(re.findall(pat, page_text)):
                            for area in page.search_for(match):
                                page.add_redact_annot(area, text=repl, fontsize=8)

                elif rtype == "redact_pattern":
                    pattern = params.get("pattern", "")
                    replacement = str(params.get("replacement", "[REDACTED]"))
                    if pattern:
                        if len(pattern) > 200:
                            print(f"[executor] PDF redact_pattern too long ({len(pattern)} chars), skipping")
                        else:
                            try:
                                compiled_pdf = re.compile(pattern)
                            except re.error as regex_err:
                                print(f"[executor] PDF redact_pattern invalid regex: {regex_err}")
                                compiled_pdf = None
                            if compiled_pdf:
                                for match in set(compiled_pdf.findall(page_text)):
                                    for area in page.search_for(str(match)):
                                        page.add_redact_annot(area, text=replacement, fontsize=8)

                elif rtype == "remove_headers_footers":
                    from collections import Counter
                    lines = page_text.splitlines()
                    counts = Counter(l.strip() for l in lines if l.strip())
                    for line, cnt in counts.items():
                        if cnt >= 2 and len(line) < 120:
                            for area in page.search_for(line):
                                page.add_redact_annot(area)

            except Exception as e:
                print(f"[executor] PDF rule {rtype} failed: {e}")

        page.apply_redactions()

    buf = io.BytesIO()
    doc.save(buf, deflate=True)
    doc.close()
    return buf.getvalue(), "application/pdf", "pdf"


def apply_transforms_docx(file_bytes: bytes, rules: list[dict]) -> tuple[bytes, str, str]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    def fix_paragraph(para):
        if not para.text.strip():
            return
        new_text = apply_document_transforms(para.text, rules)
        if new_text == para.text:
            return
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        fix_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"


def apply_document_transforms(text: str, rules: list[dict]) -> str:
    for rule in rules:
        rtype = rule["rule_type"]
        params = _parse_params_simple(rule.get("parameters", {}))
        try:
            if rtype == "strip_pii":
                text = re.sub(r'[\w.+-]+@[\w-]+\.\w+', '[EMAIL REDACTED]', text)
                text = re.sub(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', '[PHONE REDACTED]', text)
                text = re.sub(r'\b\d{3}-\d{2}-\d{4}\b', '[SSN REDACTED]', text)
                text = re.sub(r'\b(?:\d{4}[- ]?){3}\d{4}\b', '[CC REDACTED]', text)

            elif rtype == "normalize_whitespace":
                text = re.sub(r'[ \t]+', ' ', text)
                text = re.sub(r'\n{3,}', '\n\n', text)
                text = '\n'.join(l.rstrip() for l in text.splitlines())

            elif rtype == "strip_html":
                text = re.sub(r'<[^>]+>', '', text)
                for entity, char in [('&amp;','&'),('&lt;','<'),('&gt;','>'),('&nbsp;',' '),('&quot;','"')]:
                    text = text.replace(entity, char)

            elif rtype == "fix_encoding":
                replacements = {
                    'â€™': "'", 'â€œ': '"', 'â€\x9d': '"', 'â€¦': '…',
                    'â€"': '—', 'â€"': '–', 'Ã©': 'é', 'Ã¨': 'è',
                    'Ã ': 'à', 'Ã®': 'î', 'Ã´': 'ô', 'Ã¹': 'ù',
                }
                for bad, good in replacements.items():
                    text = text.replace(bad, good)

            elif rtype == "remove_blank_lines":
                lines = [l for l in text.splitlines() if l.strip()]
                text = '\n'.join(lines)

            elif rtype == "remove_headers_footers":
                from collections import Counter
                lines = text.splitlines()
                counts = Counter(l.strip() for l in lines if l.strip())
                repeated = {l for l, c in counts.items() if c >= 3 and len(l) < 120}
                text = '\n'.join(l for l in lines if l.strip() not in repeated)

            elif rtype == "redact_pattern":
                pattern = params.get("pattern", "")
                replacement = params.get("replacement", "[REDACTED]")
                if pattern:
                    # ReDoS protection: validate regex compiles and cap length
                    if len(pattern) > 200:
                        print(f"[executor] redact_pattern too long ({len(pattern)} chars), skipping")
                    else:
                        try:
                            compiled = re.compile(pattern)
                            text = compiled.sub(str(replacement), text)
                        except re.error as regex_err:
                            print(f"[executor] invalid redact_pattern regex: {regex_err}")

            elif rtype == "ner_redact":
                entities   = params.get("entities", ["PERSON", "ORG", "GPE", "DATE"])
                repl_token = str(params.get("replacement", "[REDACTED]"))
                if "PERSON" in entities:
                    text = re.sub(r'\b(?:Mr|Mrs|Ms|Dr|Prof)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b', repl_token, text)
                    text = re.sub(r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b', repl_token, text)
                if "ORG" in entities:
                    text = re.sub(r'\b[A-Z][A-Za-z\s&,\.]{2,50}(?:Inc|LLC|Ltd|LLP|Corp|Co|Company|Group|Holdings|Technologies|Solutions|Services|Associates|Consulting|Industries|Enterprises)\.?\b', repl_token, text)
                if "GPE" in entities:
                    text = re.sub(r'\b\d{1,5}\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3}(?:\s+(?:St|Ave|Blvd|Rd|Dr|Ln|Way|Ct|Pl|Terrace|Circle|Drive|Street|Avenue|Road|Lane|Court|Place)\.?)?\b', repl_token, text)
                    text = re.sub(r'\b\d{5}(?:-\d{4})?\b', repl_token, text)
                if "DATE" in entities:
                    text = re.sub(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', repl_token, text)
                    text = re.sub(r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b', repl_token, text)
                if "IP" in entities:
                    text = re.sub(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', repl_token, text)

        except Exception as e:
            print(f"[executor] skipping doc rule {rtype}: {e}")

    return text


def _parse_params_simple(raw) -> dict:
    if isinstance(raw, str):
        try:
            return json.loads(raw)
        except Exception:
            return {}
    return raw or {}


def get_db_conn():
    secret = json.loads(
        secrets.get_secret_value(SecretId=os.environ["DB_SECRET_ARN"])["SecretString"]
    )
    host = secret["host"]
    port = secret.get("port", 5432)
    user = secret["username"]
    dbname = secret.get("dbname", "cleanstack")
    rds = boto3.client("rds", region_name=os.environ.get("AWS_REGION", "us-east-1"))
    token = rds.generate_db_auth_token(DBHostname=host, Port=port, DBUsername=user)
    return psycopg2.connect(host=host, port=port, user=user, password=token, dbname=dbname, sslmode="require")


def load_raw_dataframe(file_bytes: bytes, fmt: str) -> pd.DataFrame:
    buf = io.BytesIO(file_bytes)

    if fmt in ("csv", "txt"):
        sample = file_bytes[:4096].decode("utf-8", errors="replace")
        sep = "\t" if sample.count("\t") > sample.count(",") else ","
        df = pd.read_csv(io.BytesIO(file_bytes), sep=sep, low_memory=False)
    elif fmt == "tsv":
        df = pd.read_csv(buf, sep="\t", low_memory=False)
    elif fmt in ("json", "jsonl"):
        text = file_bytes.decode("utf-8", errors="replace").strip()
        if fmt == "jsonl":
            try:
                df = pd.read_json(io.BytesIO(file_bytes), lines=True)
                return _strip_sidecar_cols(df)
            except Exception:
                pass
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                df = pd.json_normalize(parsed)
            elif isinstance(parsed, dict):
                for v in parsed.values():
                    if isinstance(v, list):
                        df = pd.json_normalize(v)
                        return _strip_sidecar_cols(df)
                df = pd.json_normalize([parsed])
            else:
                df = pd.read_json(io.BytesIO(file_bytes), lines=True)
        except Exception:
            df = pd.read_json(io.BytesIO(file_bytes), lines=True)
    elif fmt in ("xlsx", "xls"):
        xl = pd.ExcelFile(buf)
        df = xl.parse(xl.sheet_names[0])
    elif fmt == "xml":
        from lxml import etree
        root = etree.fromstring(file_bytes)
        rows = [{child.tag: child.text for child in elem} for elem in root]
        if not rows:
            rows = [{sub.tag: sub.text for sub in root}]
        df = pd.DataFrame(rows)
    elif fmt == "parquet":
        df = pd.read_parquet(buf)
    else:
        raise ValueError(f"Unsupported format for executor: {fmt}")

    return _strip_sidecar_cols(df)


def _strip_sidecar_cols(df: pd.DataFrame) -> pd.DataFrame:
    """Remove __orig_* audit sidecar columns from re-ingested processed files."""
    sidecar_cols = [c for c in df.columns if str(c).startswith("__orig_")]
    if sidecar_cols:
        print(f"[executor] stripping {len(sidecar_cols)} __orig_* sidecar columns from re-ingested file")
        df = df.drop(columns=sidecar_cols)
    return df


def _parse_params(raw) -> dict:
    if isinstance(raw, str):
        try:
            return json.loads(raw)
        except Exception:
            return {}
    return raw or {}


def _to_numeric_clean(series: pd.Series) -> pd.Series:
    """Strip currency symbols/commas then coerce to numeric."""
    cleaned = series.astype(str).str.replace(r"[$,\s]", "", regex=True)
    return pd.to_numeric(cleaned, errors="coerce")


def apply_transforms(df: pd.DataFrame, rules: list[dict]) -> pd.DataFrame:
    for rule in rules:
        rtype = rule["rule_type"]
        col = rule.get("column_name")
        params = _parse_params(rule.get("parameters"))

        try:
            if rtype == "drop_nulls":
                df_before_rule = df.copy()
                if col and col in df.columns:
                    threshold = params.get("threshold", 0.0)
                    if isinstance(threshold, (int, float)) and float(threshold) < 1.0:
                        min_count = int(len(df) * (1 - float(threshold)))
                        df = df.dropna(subset=[col], thresh=min_count)
                    else:
                        df = df.dropna(subset=[col])
                elif not col:
                    df = df.dropna()
                loss_pct = 1 - len(df) / max(len(df_before_rule), 1)
                if loss_pct > 0.20:
                    print(f"[executor] drop_nulls row-loss guard: {loss_pct:.0%} loss, reverting")
                    df = df_before_rule

            elif rtype == "deduplicate":
                subset = [col] if col and col in df.columns else None
                df = df.drop_duplicates(subset=subset, keep="first")

            elif rtype == "semantic_deduplicate":
                target_col = col if col and col in df.columns else None
                if target_col is None:
                    obj_cols = df.select_dtypes(include="object").columns.tolist()
                    target_col = obj_cols[0] if obj_cols else None
                if target_col:
                    threshold = float(params.get("threshold", 0.8))
                    num_perm  = int(params.get("num_perm", 64))

                    def _minhash_sig(text: str, n: int) -> list:
                        tokens = set(text.lower().split()) or {""}
                        return [min((hash((seed, t)) & 0x7FFFFFFF) for t in tokens) for seed in range(n)]

                    texts = df[target_col].astype(str).tolist()
                    sigs  = [_minhash_sig(t, num_perm) for t in texts]
                    keep  = []
                    kept_sigs: list = []
                    for i, sig in enumerate(sigs):
                        is_dup = any(
                            sum(a == b for a, b in zip(sig, ks)) / num_perm >= threshold
                            for ks in kept_sigs
                        )
                        if not is_dup:
                            keep.append(i)
                            kept_sigs.append(sig)
                    df = df.iloc[keep].reset_index(drop=True)

            elif rtype == "type_cast":
                if col and col in df.columns:
                    target = params.get("target_type", "str")
                    # Sidecar for irreversible casts — non-parseable originals become NaN/NaT with no recovery
                    if target in ("float", "float64", "numeric", "number", "int", "int64", "datetime", "date", "timestamp"):
                        sidecar = f"__orig_{col}"
                        if sidecar not in df.columns:
                            df[sidecar] = df[col]
                    if target in ("float", "float64", "numeric", "number"):
                        df[col] = _to_numeric_clean(df[col])
                    elif target in ("int", "int64"):
                        df[col] = _to_numeric_clean(df[col]).astype("Int64")
                    elif target == "str":
                        df[col] = df[col].astype(str)
                    elif target in ("datetime", "date", "timestamp"):
                        parsed = pd.to_datetime(df[col], infer_datetime_format=True, errors="coerce")
                        # Replace NaT with None (null) not string "NaT"
                        df[col] = parsed.dt.strftime("%Y-%m-%d").where(parsed.notna(), other=None)
                    else:
                        df[col] = df[col].astype(target, errors="ignore")

            elif rtype == "rename":
                if col and col in df.columns:
                    new_name = params.get("new_name")
                    if new_name:
                        df = df.rename(columns={col: new_name})

            elif rtype == "filter":
                if col and col in df.columns:
                    df_before_rule = df.copy()
                    operator = params.get("operator", "notnull")
                    value = params.get("value")
                    if operator == "notnull":
                        df = df[df[col].notna()]
                    elif operator == "eq":
                        df = df[df[col] == value]
                    elif operator == "neq":
                        df = df[df[col] != value]
                    elif operator == "gt":
                        df = df[_to_numeric_clean(df[col]) > float(value)]
                    elif operator == "lt":
                        df = df[_to_numeric_clean(df[col]) < float(value)]
                    loss_pct = 1 - len(df) / max(len(df_before_rule), 1)
                    if loss_pct > 0.20:
                        print(f"[executor] filter row-loss guard: {loss_pct:.0%} loss, reverting")
                        df = df_before_rule

            elif rtype == "normalize":
                if col and col in df.columns:
                    if df[col].dtype == object:
                        # Try mixed-format date parsing (pandas 2.0+)
                        try:
                            parsed = pd.to_datetime(df[col], format="mixed", dayfirst=False, errors="coerce")
                        except Exception:
                            parsed = pd.to_datetime(df[col], infer_datetime_format=True, errors="coerce")
                        # For any still-NaT values, retry with dayfirst=True
                        mask = parsed.isna() & df[col].notna() & (df[col].astype(str) != "nan")
                        if mask.any():
                            retry = pd.to_datetime(df[col][mask], format="mixed", dayfirst=True, errors="coerce")
                            parsed[mask] = retry
                        if parsed.notna().sum() > len(df) * 0.3:
                            df[col] = parsed.dt.strftime("%Y-%m-%d")
                        else:
                            df[col] = df[col].astype(str).str.strip().str.lower()
                    else:
                        numeric = pd.to_numeric(df[col], errors="coerce")
                        col_min, col_max = numeric.min(), numeric.max()
                        if col_max > col_min:
                            df[col] = (numeric - col_min) / (col_max - col_min)

            elif rtype == "fill_nulls":
                if col and col in df.columns:
                    strategy = params.get("strategy", "value")
                    fill_value = params.get("value", "Uncategorized")
                    # Sidecar for SYNTHETIC strategies — preserves originals for audit/rollback
                    if strategy in ("mean", "median", "mode"):
                        sidecar = f"__orig_{col}"
                        if sidecar not in df.columns:
                            df[sidecar] = df[col]
                    if strategy == "mean":
                        df[col] = df[col].fillna(_to_numeric_clean(df[col]).mean())
                    elif strategy == "median":
                        df[col] = df[col].fillna(_to_numeric_clean(df[col]).median())
                    elif strategy == "mode":
                        mode = df[col].mode()
                        df[col] = df[col].fillna(mode[0] if len(mode) > 0 else fill_value)
                    else:
                        df[col] = df[col].fillna(fill_value)

            elif rtype == "trim_whitespace":
                targets = [col] if (col and col in df.columns) else df.select_dtypes(include="object").columns.tolist()
                for c in targets:
                    was_null = df[c].isna()
                    df[c] = df[c].astype(str).str.strip()
                    # Null only cells that were already null (serialized to "nan" by astype(str)) or became empty after strip
                    # Preserves legitimate string value "nan" (e.g. currency ticker, city name)
                    df[c] = df[c].where(~(was_null | (df[c] == "")), other=None)

            elif rtype == "ffill":
                if col and col in df.columns:
                    sidecar = f"__orig_{col}"
                    if sidecar not in df.columns:
                        df[sidecar] = df[col]
                    df[col] = df[col].ffill()

            elif rtype == "bfill":
                if col and col in df.columns:
                    sidecar = f"__orig_{col}"
                    if sidecar not in df.columns:
                        df[sidecar] = df[col]
                    df[col] = df[col].bfill()

            elif rtype == "bool_cast":
                if col and col in df.columns:
                    TRUE_VALS = {"true", "yes", "1", "t", "y", "on"}
                    FALSE_VALS = {"false", "no", "0", "f", "n", "off"}
                    target = params.get("target", "bool")
                    def _cast_bool(v):
                        if pd.isna(v):
                            return v
                        s = str(v).strip().lower()
                        if s in TRUE_VALS:
                            return True if target == "bool" else (1 if target == "int" else "true")
                        if s in FALSE_VALS:
                            return False if target == "bool" else (0 if target == "int" else "false")
                        return v  # unknown value → keep original unchanged
                    df[col] = df[col].apply(_cast_bool)

            elif rtype == "outlier_cap":
                if col and col in df.columns:
                    multiplier = float(params.get("iqr_multiplier", 1.5))
                    numeric = pd.to_numeric(df[col], errors="coerce")
                    q1 = numeric.quantile(0.25)
                    q3 = numeric.quantile(0.75)
                    iqr = q3 - q1
                    lower = q1 - multiplier * iqr
                    upper = q3 + multiplier * iqr
                    sidecar = f"__orig_{col}"
                    if sidecar not in df.columns:
                        df[sidecar] = df[col]
                    df[col] = numeric.clip(lower=lower, upper=upper)

            elif rtype == "multi_currency_strip":
                targets = [col] if (col and col in df.columns) else df.select_dtypes(include="object").columns.tolist()
                CURRENCY_PATTERN = r'[$€£¥₹₩]|(?:USD|EUR|GBP|JPY|INR|CAD|AUD|CHF)\s*'
                for c in targets:
                    if df[c].dtype != object:
                        continue
                    cleaned = df[c].astype(str).str.replace(CURRENCY_PATTERN, "", regex=True)
                    cleaned = cleaned.str.replace(",", "", regex=False).str.strip()
                    numeric = pd.to_numeric(cleaned, errors="coerce")
                    non_null = df[c].notna().sum()
                    parsed_ok = numeric.notna().sum()
                    if non_null > 0 and parsed_ok / non_null > 0.5:
                        sidecar = f"__orig_{c}"
                        if sidecar not in df.columns:
                            df[sidecar] = df[c]
                        df[c] = numeric
                    else:
                        print(f"[executor] multi_currency_strip: {c} — only {parsed_ok}/{non_null} parsed, skipping")

            elif rtype == "filter_extended":
                if col and col in df.columns:
                    operator = params.get("operator", "notnull")
                    value = params.get("value")
                    values = params.get("values", [])
                    pattern = params.get("pattern", "")
                    df_before_rule = df.copy()
                    if operator == "notnull":
                        df = df[df[col].notna()]
                    elif operator == "eq":
                        df = df[df[col] == value]
                    elif operator == "neq":
                        df = df[df[col] != value]
                    elif operator == "gt":
                        df = df[_to_numeric_clean(df[col]) > float(value)]
                    elif operator == "lt":
                        df = df[_to_numeric_clean(df[col]) < float(value)]
                    elif operator == "gte":
                        df = df[_to_numeric_clean(df[col]) >= float(value)]
                    elif operator == "lte":
                        df = df[_to_numeric_clean(df[col]) <= float(value)]
                    elif operator == "contains":
                        df = df[df[col].astype(str).str.contains(str(value), na=False)]
                    elif operator == "not_contains":
                        df = df[~df[col].astype(str).str.contains(str(value), na=False)]
                    elif operator == "in":
                        df = df[df[col].isin(values)]
                    elif operator == "not_in":
                        df = df[~df[col].isin(values)]
                    elif operator == "regex":
                        if len(pattern) <= 200:
                            try:
                                df = df[df[col].astype(str).str.match(pattern, na=False)]
                            except re.error as e:
                                print(f"[executor] filter_extended regex error: {e}")
                    elif operator == "startswith":
                        df = df[df[col].astype(str).str.startswith(str(value), na=False)]
                    elif operator == "endswith":
                        df = df[df[col].astype(str).str.endswith(str(value), na=False)]
                    # Per-rule row loss guard (Safeguard 3)
                    loss_pct = (len(df_before_rule) - len(df)) / max(len(df_before_rule), 1)
                    if loss_pct > 0.20:
                        print(f"[executor] SAFETY: filter_extended on {col} deleted {loss_pct:.1%} rows — restoring")
                        df = df_before_rule

            elif rtype == "split_column":
                if col and col in df.columns:
                    delimiter = params.get("delimiter", "|")
                    new_col_names = params.get("new_columns", [])
                    max_splits = int(params.get("max_splits", -1))
                    split_df = df[col].astype(str).str.split(
                        pat=re.escape(delimiter),
                        n=max_splits if max_splits > 0 else -1,
                        expand=True,
                    )
                    for i in range(split_df.shape[1]):
                        new_name = new_col_names[i] if i < len(new_col_names) else f"{col}_part{i + 1}"
                        df[new_name] = split_df[i]
                    # Source column preserved — user can drop manually

            elif rtype == "column_header_normalize":
                def _to_snake(name: str) -> str:
                    s = str(name).strip()
                    s = re.sub(r'[^\w\s]', '_', s)
                    s = re.sub(r'\s+', '_', s)
                    s = re.sub(r'([a-z])([A-Z])', r'\1_\2', s)
                    s = re.sub(r'_+', '_', s)
                    return s.lower().strip('_')
                rename_map = {}
                for c in df.columns:
                    new_name = _to_snake(str(c))
                    if new_name != str(c):
                        rename_map[c] = new_name
                if rename_map:
                    df = df.rename(columns=rename_map)
                    print(f"[executor] column_header_normalize: renamed {len(rename_map)} columns")

            elif rtype == "ner_redact":
                entities   = params.get("entities", ["PERSON", "ORG", "GPE", "DATE"])
                repl_token = str(params.get("replacement", "[REDACTED]"))
                targets = [col] if (col and col in df.columns) else df.select_dtypes(include="object").columns.tolist()

                _NER_PATTERNS: list[tuple[str, str]] = []
                if "PERSON" in entities:
                    _NER_PATTERNS += [
                        # Title + name  (Mr. John Smith)
                        (r'\b(?:Mr|Mrs|Ms|Dr|Prof)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b', repl_token),
                        # First Last (two capitalised words, 3+ chars each, not start of sentence POS)
                        (r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b', repl_token),
                    ]
                if "ORG" in entities:
                    _NER_PATTERNS += [
                        (r'\b[A-Z][A-Za-z\s&,\.]{2,50}(?:Inc|LLC|Ltd|LLP|Corp|Co|Company|Group|Holdings|Technologies|Solutions|Services|Associates|Consulting|Industries|Enterprises)\.?\b', repl_token),
                    ]
                if "GPE" in entities:
                    _us_states = (
                        "Alabama|Alaska|Arizona|Arkansas|California|Colorado|Connecticut|Delaware|"
                        "Florida|Georgia|Hawaii|Idaho|Illinois|Indiana|Iowa|Kansas|Kentucky|"
                        "Louisiana|Maine|Maryland|Massachusetts|Michigan|Minnesota|Mississippi|"
                        "Missouri|Montana|Nebraska|Nevada|New Hampshire|New Jersey|New Mexico|"
                        "New York|North Carolina|North Dakota|Ohio|Oklahoma|Oregon|Pennsylvania|"
                        "Rhode Island|South Carolina|South Dakota|Tennessee|Texas|Utah|Vermont|"
                        "Virginia|Washington|West Virginia|Wisconsin|Wyoming"
                    )
                    _NER_PATTERNS += [
                        # US street address
                        (r'\b\d{1,5}\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3}(?:\s+(?:St|Ave|Blvd|Rd|Dr|Ln|Way|Ct|Pl|Terrace|Circle|Drive|Street|Avenue|Road|Lane|Court|Place)\.?)?\b', repl_token),
                        # US state names
                        (rf'\b(?:{_us_states})\b', repl_token),
                        # US ZIP code
                        (r'\b\d{5}(?:-\d{4})?\b', repl_token),
                    ]
                if "DATE" in entities:
                    _NER_PATTERNS += [
                        # Month DD, YYYY
                        (r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b', repl_token),
                        # MM/DD/YYYY or DD/MM/YYYY
                        (r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b', repl_token),
                    ]
                if "IP" in entities:
                    _NER_PATTERNS += [
                        (r'\b(?:\d{1,3}\.){3}\d{1,3}\b', repl_token),
                    ]

                def _apply_ner(text: str) -> str:
                    for pattern, replacement in _NER_PATTERNS:
                        text = re.sub(pattern, replacement, text)
                    return text

                for c in targets:
                    df[c] = df[c].astype(str).apply(_apply_ner)
                    df[c] = df[c].replace({"nan": None, "": None})

        except Exception as e:
            print(f"[executor] skipping rule {rtype} on {col}: {e}")

    return df


def compute_quality_profile(df: pd.DataFrame) -> dict:
    total_cells = df.size or 1
    total_rows = len(df)

    null_count = df.isnull().sum().sum()
    null_pct = round(null_count / total_cells * 100, 2)

    dup_count = df.duplicated().sum()
    dup_pct = round(dup_count / max(total_rows, 1) * 100, 2)

    type_mismatches = 0
    for col in df.columns:
        if df[col].dtype == object:
            numeric_count = pd.to_numeric(df[col], errors="coerce").notna().sum()
            if 0 < numeric_count < len(df[col]):
                type_mismatches += 1

    outlier_count = 0
    for col in df.select_dtypes(include=[np.number]).columns:
        q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
        iqr = q3 - q1
        outlier_count += int(
            df[(df[col] < q1 - 1.5 * iqr) | (df[col] > q3 + 1.5 * iqr)][col].count()
        )

    null_penalty = min(null_pct * 0.5, 30)
    dup_penalty = min(dup_pct * 0.3, 20)
    type_penalty = min(type_mismatches * 5, 20)
    outlier_penalty = min(outlier_count / max(total_rows, 1) * 100 * 0.1, 10)
    score = max(0, round(100 - null_penalty - dup_penalty - type_penalty - outlier_penalty))

    column_stats = {}
    for col in df.columns:
        series = df[col]
        stat = {
            "type": str(series.dtype),
            "null_count": int(series.isnull().sum()),
            "null_pct": round(series.isnull().mean() * 100, 2),
            "unique_count": int(series.nunique()),
            "sample_values": [str(v) for v in series.dropna().head(5).tolist()],
        }
        if pd.api.types.is_numeric_dtype(series):
            stat["min"] = float(series.min()) if not series.empty else None
            stat["max"] = float(series.max()) if not series.empty else None
        column_stats[str(col)] = stat

    return {
        "quality_score": score,
        "total_rows": total_rows,
        "null_percentage": null_pct,
        "duplicate_percentage": dup_pct,
        "type_mismatch_count": type_mismatches,
        "outlier_count": outlier_count,
        "column_stats": column_stats,
    }


def schema_hash(df: pd.DataFrame) -> tuple[str, dict]:
    col_defs = {str(col): str(df[col].dtype) for col in df.columns}
    h = hashlib.sha256(json.dumps(col_defs, sort_keys=True).encode()).hexdigest()
    return h, col_defs


def _maybe_auto_iterate(cur, conn, s3_client, run_id, pipeline_id, raw_s3_key, fmt,
                         processed_key, iteration, processed_score):
    """Check improvement vs parent and create next pass if warranted."""
    import uuid as _uuid
    try:
        # Fetch parent_run_id, then get its processed quality score
        cur.execute("SELECT parent_run_id FROM pipeline_runs WHERE id = %s", (run_id,))
        parent_row = cur.fetchone()
        if not parent_row or not parent_row[0]:
            return

        parent_run_id = parent_row[0]
        cur.execute(
            "SELECT quality_score FROM data_profiles WHERE run_id = %s AND stage = 'processed'",
            (parent_run_id,)
        )
        parent_score_row = cur.fetchone()
        if not parent_score_row:
            return

        parent_score = float(parent_score_row[0])
        if parent_score <= 0:
            return

        improvement_pct = (float(processed_score) - parent_score) / parent_score * 100
        print(f"[executor] Auto-iterate: improvement={improvement_pct:.1f}% (pass {iteration}→{iteration+1})")

        if improvement_pct < 5.0:
            print("[executor] Improvement < 5% — stopping auto-clean loop")
            return

        # Create next pass run
        raw_bucket = os.environ["S3_RAW_BUCKET"]
        processed_bucket = os.environ["S3_PROCESSED_BUCKET"]
        user_id = raw_s3_key.split("/")[0]
        new_run_id = str(_uuid.uuid4())
        new_raw_key = f"{user_id}/{pipeline_id}/{new_run_id}/raw.{fmt}"

        # Copy processed → new raw (triggers profiler via S3 event)
        s3_client.copy_object(
            Bucket=raw_bucket,
            CopySource={"Bucket": processed_bucket, "Key": processed_key},
            Key=new_raw_key,
        )

        cur.execute(
            """INSERT INTO pipeline_runs
               (id, pipeline_id, status, file_format, raw_s3_key, started_at,
                iteration, parent_run_id, auto_mode)
               VALUES (%s, %s, 'pending', %s, %s, now(), %s, %s, TRUE)""",
            (new_run_id, pipeline_id, fmt, new_raw_key, iteration + 1, run_id)
        )
        conn.commit()
        print(f"[executor] Auto-iterate: created pass {iteration+1} run {new_run_id}")

    except Exception as e:
        import traceback
        print(f"[executor] Auto-iterate error (non-fatal, run {run_id}): {e}")
        print(traceback.format_exc())
        # Don't raise — auto-iterate failure must not fail the current run


def handler(event, context):
    record = event["Records"][0]
    body = json.loads(record["body"])
    run_id = body["run_id"]

    conn = get_db_conn()
    cur = conn.cursor()

    try:
        cur.execute(
            "UPDATE pipeline_runs SET status = 'running' WHERE id = %s",
            (run_id,)
        )
        conn.commit()

        # Fetch run metadata + pipeline settings
        cur.execute(
            """SELECT pr.pipeline_id, pr.raw_s3_key, pr.file_format, pr.mode,
                      pr.iteration, pr.parent_run_id, pr.auto_mode, pr.row_count_raw,
                      p.auto_delete_raw
               FROM pipeline_runs pr
               JOIN pipelines p ON pr.pipeline_id = p.id
               WHERE pr.id = %s""",
            (run_id,)
        )
        row = cur.fetchone()
        pipeline_id, raw_s3_key, file_format, run_mode, iteration, parent_run_id, auto_mode, row_count_raw, auto_delete_raw = row
        run_mode = run_mode or "tabular"
        iteration = iteration or 1
        auto_mode = auto_mode or False
        auto_delete_raw = auto_delete_raw if auto_delete_raw is not None else True

        # Fetch approved rules ordered by index
        cur.execute(
            """SELECT rule_type, column_name, parameters
               FROM transform_rules
               WHERE run_id = %s AND status = 'approved'
               ORDER BY order_index ASC""",
            (run_id,)
        )
        all_rules = [
            {"rule_type": r[0], "column_name": r[1], "parameters": r[2]}
            for r in cur.fetchall()
        ]

        # Deduplicate document rules by rule_type (S3 at-least-once can produce duplicates).
        # For tabular rules, dedup by (rule_type, column_name) to keep per-column rules.
        seen_doc: set[str] = set()
        rules = []
        for r in all_rules:
            key = r["rule_type"] if r["column_name"] is None else f"{r['rule_type']}::{r['column_name']}"
            if key not in seen_doc:
                seen_doc.add(key)
                rules.append(r)

        # Read raw file from S3
        raw_bucket = os.environ["S3_RAW_BUCKET"]
        obj = s3.get_object(Bucket=raw_bucket, Key=raw_s3_key)
        file_bytes = obj["Body"].read()
        fmt = file_format or raw_s3_key.rsplit(".", 1)[-1].lower()

        processed_bucket = os.environ["S3_PROCESSED_BUCKET"]

        if run_mode == "document":
            text_key = "/".join(raw_s3_key.rsplit("/", 1)[:-1]) + "/extracted_text.txt"

            if fmt == "pdf":
                out_bytes, content_type, ext = apply_transforms_pdf(file_bytes, rules)
            elif fmt == "docx":
                out_bytes, content_type, ext = apply_transforms_docx(file_bytes, rules)
            else:
                # txt and other text formats — use pre-extracted text
                try:
                    text_obj = s3.get_object(Bucket=raw_bucket, Key=text_key)
                    text = text_obj["Body"].read().decode("utf-8")
                except Exception:
                    text = extract_text(file_bytes, fmt)
                text = apply_document_transforms(text, rules)
                out_bytes = text.encode("utf-8")
                content_type = "text/plain"
                ext = "txt"

            processed_key = f"processed/{pipeline_id}/{run_id}/output.{ext}"
            s3.put_object(Bucket=processed_bucket, Key=processed_key,
                          Body=out_bytes, ContentType=content_type)

            # Line count from extracted text for all doc formats
            try:
                text_obj = s3.get_object(Bucket=raw_bucket, Key=text_key)
                line_count = len(text_obj["Body"].read().decode("utf-8").splitlines())
            except Exception:
                line_count = len(out_bytes.decode("utf-8", errors="replace").splitlines())

            profile = {
                "quality_score": 95,
                "total_rows": line_count,
                "null_percentage": 0.0,
                "duplicate_percentage": 0.0,
                "type_mismatch_count": 0,
                "outlier_count": 0,
                "column_stats": {},
            }
        else:
            df = load_raw_dataframe(file_bytes, fmt)
            input_row_count = len(df)
            df = apply_transforms(df, rules)

            # Row count guard for auto-mode passes 2+ — abort if >10% rows deleted
            if auto_mode and iteration > 1 and row_count_raw:
                output_row_count = len(df)
                loss_pct = (input_row_count - output_row_count) / max(input_row_count, 1)
                if loss_pct > 0.10:
                    print(f"[executor] Row count guard triggered: {loss_pct:.1%} loss — aborting, keeping parent output")
                    cur.execute(
                        "UPDATE pipeline_runs SET status = 'failed', error_message = %s WHERE id = %s",
                        (f"Auto-clean aborted: {loss_pct:.1%} row loss exceeds 10% safety threshold", run_id),
                    )
                    conn.commit()
                    return {"statusCode": 200, "run_id": run_id, "aborted": True}

            # Write processed file in native format to S3
            file_bytes, content_type, ext = save_dataframe(df, fmt)
            processed_key = f"processed/{pipeline_id}/{run_id}/output.{ext}"
            s3.put_object(
                Bucket=processed_bucket,
                Key=processed_key,
                Body=file_bytes,
                ContentType=content_type,
            )
            # Strip __orig_* sidecar columns before profiling — they inflate type_mismatch and trigger false drift alerts
            df_for_profile = df[[c for c in df.columns if not str(c).startswith("__orig_")]]
            profile = compute_quality_profile(df_for_profile)

        cur.execute(
            """INSERT INTO data_profiles
               (run_id, stage, quality_score, total_rows, null_percentage,
                duplicate_percentage, type_mismatch_count, outlier_count, column_stats)
               VALUES (%s, 'processed', %s, %s, %s, %s, %s, %s, %s)""",
            (
                run_id,
                float(profile["quality_score"]),
                int(profile["total_rows"]),
                float(profile["null_percentage"]),
                float(profile["duplicate_percentage"]),
                int(profile["type_mismatch_count"]),
                int(profile["outlier_count"]),
                json.dumps(_sanitize_nan(profile["column_stats"]), cls=_NpEncoder),
            ),
        )

        cur.execute(
            """UPDATE pipeline_runs
               SET status = 'completed',
                   processed_s3_key = %s,
                   row_count_processed = %s,
                   completed_at = now()
               WHERE id = %s""",
            (processed_key, profile["total_rows"], run_id),
        )
        conn.commit()

        # Schema drift detection — tabular only
        if run_mode == "document":
            # Auto-iterate for document mode — copy raw file BEFORE deleting it
            if auto_mode and iteration < 3:
                _maybe_auto_iterate(
                    cur, conn, s3, run_id, pipeline_id, raw_s3_key, fmt,
                    processed_key, iteration, profile["quality_score"]
                )
            # Delete raw file AFTER auto-iterate has copied it (if applicable)
            if auto_delete_raw:
                try:
                    raw_bucket = os.environ["S3_RAW_BUCKET"]
                    s3.delete_object(Bucket=raw_bucket, Key=raw_s3_key)
                    print(f"[executor] deleted raw file s3://{raw_bucket}/{raw_s3_key}")
                except Exception as del_err:
                    print(f"[executor] raw file deletion failed (non-fatal): {del_err}")
            return {"statusCode": 200, "run_id": run_id}

        # Strip __orig_* sidecar columns before schema hash — sidecars must not trigger drift alerts
        df_for_schema = df[[c for c in df.columns if not str(c).startswith("__orig_")]]
        new_hash, col_defs = schema_hash(df_for_schema)

        cur.execute(
            """SELECT schema_hash FROM schema_snapshots
               WHERE pipeline_id = %s
               ORDER BY created_at DESC LIMIT 1""",
            (pipeline_id,)
        )
        last = cur.fetchone()

        cur.execute(
            """INSERT INTO schema_snapshots (pipeline_id, run_id, schema_hash, column_definitions)
               VALUES (%s, %s, %s, %s)""",
            (pipeline_id, run_id, new_hash, json.dumps(col_defs))
        )
        conn.commit()

        if last and last[0] != new_hash:
            sns_topic = os.environ.get("SNS_DRIFT_TOPIC_ARN")
            if sns_topic:
                sns.publish(
                    TopicArn=sns_topic,
                    Subject=f"Schema drift detected — pipeline {pipeline_id}",
                    Message=json.dumps({
                        "pipeline_id": pipeline_id,
                        "run_id": run_id,
                        "previous_hash": last[0],
                        "new_hash": new_hash,
                        "new_schema": col_defs,
                    }),
                )

        # Auto-iterate for tabular mode — copy raw file BEFORE deleting it
        if auto_mode and iteration < 3:
            _maybe_auto_iterate(
                cur, conn, s3, run_id, pipeline_id, raw_s3_key, fmt,
                processed_key, iteration, profile["quality_score"]
            )

        # Delete raw file AFTER auto-iterate has copied it (if applicable)
        if auto_delete_raw:
            try:
                raw_bucket = os.environ["S3_RAW_BUCKET"]
                s3.delete_object(Bucket=raw_bucket, Key=raw_s3_key)
                print(f"[executor] deleted raw file s3://{raw_bucket}/{raw_s3_key}")
            except Exception as del_err:
                print(f"[executor] raw file deletion failed (non-fatal): {del_err}")

    except Exception as e:
        conn.rollback()
        cur.execute(
            "UPDATE pipeline_runs SET status = 'failed', error_message = %s WHERE id = %s",
            (str(e), run_id),
        )
        conn.commit()
        raise
    finally:
        cur.close()
        conn.close()

    return {"statusCode": 200, "run_id": run_id}
